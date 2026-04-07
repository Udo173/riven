# charter_generator.py
# PM Tool - Project Charter Generator
# Claude Architecture | OpenCode Implementation
# Path: C:\Users\USass\Billi_Bilder\PM_Tool\core\

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


def create_project_charter(data: dict, output_path: str = None):
    """
    Erstellt ein professionelles Project Charter Dokument.

    Args:
        data: Dictionary mit Projekt-Daten
        output_path: Speicherpfad (optional)

    Returns:
        str: Pfad zur gespeicherten Datei
    """
    doc = Document()

    # === SEITENRÄNDER ===
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === HELPER FUNCTIONS ===
    def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = color
        return h

    def add_field(label, value, bold_label=True):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = bold_label
        run_label.font.size = Pt(11)
        run_value = p.add_run(str(value))
        run_value.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_table_row(table, col1, col2, header=False):
        row = table.add_row()
        for i, text in enumerate([col1, col2]):
            cell = row.cells[i]
            cell.text = text
            run = cell.paragraphs[0].runs[0]
            run.bold = header
            run.font.size = Pt(10)
        return row

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT CHARTER")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(data.get("project_name", "Projektname"))
    sub_run.bold = True
    sub_run.font.size = Pt(16)

    doc.add_paragraph()

    # === 1. PROJEKT-INFO ===
    add_heading("1. Projekt-Informationen")
    add_field("Projektname", data.get("project_name", "-"))
    add_field("Projektleiter", data.get("project_manager", "-"))
    add_field("Auftraggeber", data.get("sponsor", "-"))
    add_field("Startdatum", data.get("start_date", str(date.today())))
    add_field("Enddatum", data.get("end_date", "-"))
    add_field("Version", data.get("version", "1.0"))
    add_field("Status", data.get("status", "In Planung"))

    # === 2. PROJEKTZIEL ===
    add_heading("2. Projektziel")
    doc.add_paragraph(data.get("objective", "Ziel noch nicht definiert."))

    # === 3. SCOPE ===
    add_heading("3. Projektumfang (Scope)")

    doc.add_paragraph("Im Scope:", style="List Bullet").runs[0].bold = True
    for item in data.get("in_scope", ["Noch nicht definiert"]):
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()
    doc.add_paragraph("Ausserhalb des Scope:", style="List Bullet").runs[0].bold = True
    for item in data.get("out_of_scope", ["Noch nicht definiert"]):
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)

    # === 4. STAKEHOLDER ===
    add_heading("4. Stakeholder")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Name", "Rolle", "Interesse"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for sh in data.get("stakeholders", [{"name": "-", "role": "-", "interest": "-"}]):
        row = table.add_row().cells
        row[0].text = sh.get("name", "-")
        row[1].text = sh.get("role", "-")
        row[2].text = sh.get("interest", "-")

    # === 5. MEILENSTEINE ===
    doc.add_paragraph()
    add_heading("5. Meilensteine")
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(["Meilenstein", "Datum", "Status"]):
        hdr2[i].text = h
        hdr2[i].paragraphs[0].runs[0].bold = True

    for ms in data.get("milestones", [{"name": "-", "date": "-", "status": "Offen"}]):
        row = table2.add_row().cells
        row[0].text = ms.get("name", "-")
        row[1].text = ms.get("date", "-")
        row[2].text = ms.get("status", "Offen")

    # === 6. RISIKEN ===
    doc.add_paragraph()
    add_heading("6. Risiken")
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Light Grid Accent 1"
    hdr3 = table3.rows[0].cells
    for i, h in enumerate(["Risiko", "Wahrscheinlichkeit", "Massnahme"]):
        hdr3[i].text = h
        hdr3[i].paragraphs[0].runs[0].bold = True

    for r in data.get("risks", [{"risk": "-", "probability": "-", "action": "-"}]):
        row = table3.add_row().cells
        row[0].text = r.get("risk", "-")
        row[1].text = r.get("probability", "-")
        row[2].text = r.get("action", "-")

    # === 7. BUDGET ===
    doc.add_paragraph()
    add_heading("7. Budget")
    add_field("Geplantes Budget", data.get("budget", "Noch nicht definiert"))
    add_field("Budget-Verantwortlicher", data.get("budget_owner", "-"))

    # === 8. GENEHMIGUNG ===
    doc.add_paragraph()
    add_heading("8. Genehmigung")
    doc.add_paragraph(
        f"Erstellt am: {date.today().strftime('%d.%m.%Y')} | "
        f"Erstellt von: {data.get('project_manager', '-')}"
    )
    doc.add_paragraph("\n_______________________          _______________________")
    doc.add_paragraph("Projektleiter                              Auftraggeber")

    # === SPEICHERN ===
    if not output_path:
        project_safe = data.get("project_name", "Projekt").replace(" ", "_")
        output_path = (
            f"C:\\Users\\USass\\Billi_Bilder\\PM_Tool\\{project_safe}_Charter.docx"
        )

    doc.save(output_path)
    print(f"[OK] Charter gespeichert: {output_path}")
    return output_path


# === TEST-RUN ===
if __name__ == "__main__":
    test_data = {
        "project_name": "PM Tool",
        "project_manager": "Udo Sassik",
        "sponsor": "Billi_Bilder",
        "start_date": "07.04.2026",
        "end_date": "30.06.2026",
        "version": "1.0",
        "status": "In Planung",
        "objective": (
            "Entwicklung eines PM Tools zur automatischen Erstellung von "
            "Project Charters, Canvas Diagrammen und Gantt Charts mit "
            "integriertem KI-Assistenten (Ollama llama3.1)."
        ),
        "in_scope": [
            "Project Charter Generator (python-docx)",
            "Canvas Diagramm Generator (python-pptx)",
            "Gantt Chart Generator (openpyxl)",
            "KI-Assistent Integration (Ollama)",
        ],
        "out_of_scope": [
            "Web-Interface (Phase 2)",
            "Cloud-Deployment",
        ],
        "stakeholders": [
            {
                "name": "Udo Sassik",
                "role": "Projektleiter",
                "interest": "Vollständige Kontrolle",
            },
            {"name": "OpenCode", "role": "Entwicklung", "interest": "Code-Qualität"},
            {
                "name": "Claude",
                "role": "Architektur",
                "interest": "Strategie & Planung",
            },
        ],
        "milestones": [
            {
                "name": "Charter Generator fertig",
                "date": "10.04.2026",
                "status": "In Arbeit",
            },
            {
                "name": "Canvas Generator fertig",
                "date": "15.04.2026",
                "status": "Offen",
            },
            {"name": "Gantt Generator fertig", "date": "20.04.2026", "status": "Offen"},
            {
                "name": "KI-Assistent integriert",
                "date": "30.04.2026",
                "status": "Offen",
            },
        ],
        "risks": [
            {
                "risk": "python-pptx Kompatibilität",
                "probability": "Niedrig",
                "action": "Version fixieren",
            },
            {
                "risk": "Ollama API Aenderungen",
                "probability": "Mittel",
                "action": "Lokale Version pinnen",
            },
        ],
        "budget": "0 Euro (nur kostenlose Tools)",
        "budget_owner": "Udo Sassik",
    }

    create_project_charter(test_data)
